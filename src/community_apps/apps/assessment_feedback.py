# assessment_feedback.py

import streamlit as st
import pandas as pd
from langchain import hub
from langchain_openai import AzureChatOpenAI, ChatOpenAI
import docx
import io
import asyncio
import os
import numpy as np  # Add numpy import for random selection
from textwrap import fill


st.title("Assessment Feedback")

# Initialize all session state variables
if "column_config" not in st.session_state:
    st.session_state.column_config = None
if "selected_columns" not in st.session_state:
    st.session_state.selected_columns = None
if "test_size" not in st.session_state:
    st.session_state.test_size = None
if "df" not in st.session_state:
    st.session_state.df = None
if "processed_df" not in st.session_state:
    st.session_state.processed_df = None
if "full_results_df" not in st.session_state:
    st.session_state.full_results_df = None
if "instructions" not in st.session_state:
    st.session_state.instructions = ""
if "previous_selection" not in st.session_state:
    st.session_state.previous_selection = []

# Default instructions from instructions.md
DEFAULT_INSTRUCTIONS = """# L6 Zombie Fire Explanation

Summative To study changing fires in the Arctic, scientists took pictures of fires from space, using special cameras. The first image (top-left), acquired in September 2015, shows the burn scar from the Soda Creek Fire, which scorched nearly 17,000 acres in southwest Alaska near the Kuskokwim River. The fire was never completely extinguished before winter set in. In April 2016 (top-right), the fire continued to smolder in the peat under a layer of snow. When the snow finally melted in late May (bottom-left), the additional heat and oxygen caused flames to re-emerge and quickly spread.

| Instructions: Use evidence from your investigations in class to explain both questions below:    (A) How was there enough matter and energy in the system for a zombie fire to burn under ice?    (B) What will happen in the future if temperatures continue to rise? |  |
| :---- | :---- |
| **Sentence starters:**    The components of the zombie fire system are…    Zombie fires can burn under ice because…     The peat formed by the process of….    If the temperatures of the Earth continue to rise…  | Energy flows through the system from… to… through the process of…   The matter in the system transforms from… to … through the process of…  In the past, the arctic was different than today because….  |
| **Key vocabulary to include:**    Photosynthesis    Cellular respiration / Decomposition | Oxygen Temperature Energy |

| 4 Extending | 3 Proficient | 2 Approaching | 1 Beginning |
| :---- | :---- | :---- | :---- |
| Constructs an explanation about zombie fires based on empirical evidence and makes specific connections to multiple disciplinary ideas. Cites specific evidence from multiple class activities.   | Constructs an explanation about zombie fires based on empirical evidence and makes a connection to multiple disciplinary ideas. Cites specific evidence from class.  (Don't use words like more or a lot.) | Constructs an explanation about zombie fires based on empirical evidence and begins to make a connection to multiple disciplinary ideas. References evidence from class (without citing specific data).  | Constructs an explanation (**including interactions**) to describe zombie fires with **minimal** reference to empirical evidence.  |

| Check | Attribute |
| ----- | ----- |
|  | Did I make a claim about matter and energy in the zombie fire system?  Does my claim discuss energy both flowing into and out of the zombie fire system?  |
|  | Did I include specific evidence from activities we did in class? Burning fuel samples demonstration Yeast/Decomposition Lab Elodea/Earth's Tilt lab Progress Tracker |
|  | Did I explain how my evidence supports my claim?  Did I apply multiple scientific ideas to my explanation? Carbon cycle Cellular respiration Photosynthesis Decomposition Solar radiation |
|  | Did I make a prediction about the future?  |
|  | Did I reference evidence from this unit to support my prediction? |
|  | Is my writing legible?  Is my writing concise? |"""


def prepare_batch(df, selected_columns, start_idx=0, batch_size=None):
    """Prepare a batch of samples for processing."""
    if not selected_columns:
        raise ValueError("No columns selected for student work")

    if not isinstance(df, pd.DataFrame):
        raise ValueError("Invalid dataframe provided")

    if batch_size is None:
        batch_size = len(df) - start_idx

    end_idx = min(start_idx + batch_size, len(df))
    batch_df = df.iloc[start_idx:end_idx].copy()

    # Validate selected columns exist
    missing_cols = [col for col in selected_columns if col not in batch_df.columns]
    if missing_cols:
        raise ValueError(f"Selected columns not found: {', '.join(missing_cols)}")

    instructions = st.session_state.get("instructions", "")

    # Prepare inputs for the batch with better error handling
    inputs = []
    valid_indices = []
    skipped_count = 0

    for idx, row in batch_df.iterrows():
        # Concatenate content from all selected columns
        student_work_parts = []
        for col in selected_columns:
            content = row[col]
            if not pd.isna(content) and str(content).strip():
                student_work_parts.append(f"{col}: {str(content).strip()}")

        # Skip if all parts are empty
        if not student_work_parts:
            skipped_count += 1
            continue

        # Join parts with newlines
        student_response = "\n\n".join(student_work_parts)

        inputs.append(
            {
                "student_response": student_response,
                "instructions": instructions,
            }
        )
        valid_indices.append(idx)

    if skipped_count > 0:
        st.warning(f"Skipped {skipped_count} empty responses")

    if not inputs:
        raise ValueError("No valid student work found in selected columns")

    # Return only the rows with valid responses
    return batch_df.loc[valid_indices], inputs


async def process_batch(chain, batch_df, inputs):
    """Process a batch of samples and return results."""
    try:
        # Process the batch
        results = await chain.abatch(inputs)

        # Get all unique keys from the results excluding input fields
        result_keys = set()
        for result in results:
            result_keys.update(result.keys())
        result_keys = result_keys - {"instructions", "student_response"}

        # Dynamically add result columns to the dataframe
        for key in result_keys:
            batch_df[key] = [r.get(key, "") for r in results]

        return batch_df, results, True

    except Exception as e:
        st.error(f"Error processing batch: {str(e)}")
        return batch_df, None, False


def get_chain():
    """Initialize and return the LangChain chain."""
    if "chain" not in st.session_state:
        prompt = hub.pull("hey-aw/generate_student_feedback")

        # Get model name from secrets
        MODEL_NAME = st.secrets.get("OPENAI_MODEL", "gpt-4")

        # Try Azure OpenAI first, fall back to OpenAI
        azure_key = st.secrets.get("AZURE_OPENAI_API_KEY")
        if azure_key:
            model = AzureChatOpenAI(
                api_key=azure_key,
                azure_endpoint=st.secrets.get("AZURE_OPENAI_ENDPOINT"),
                azure_deployment=st.secrets.get("AZURE_OPENAI_DEPLOYMENT"),
                model_name=MODEL_NAME,
                temperature=0,  # Add temperature=0 for more consistent results
            )
        else:
            openai_key = st.secrets.get("OPENAI_API_KEY")
            if not openai_key:
                st.error(
                    "No OpenAI API key found. Please set either AZURE_OPENAI_API_KEY or OPENAI_API_KEY in secrets."
                )
                return None
            model = ChatOpenAI(
                api_key=openai_key,
                model_name=MODEL_NAME,
                temperature=0,  # Add temperature=0 for more consistent results
            )

        st.session_state.chain = prompt | model
    return st.session_state.chain


async def generate_feedback(df, selected_columns):
    """Main feedback generation function."""
    chain = get_chain()
    if chain is None:
        return None, None

    batch_df, inputs = prepare_batch(
        df, selected_columns, batch_size=st.session_state.test_size
    )
    processed_df, results, success = await process_batch(chain, batch_df, inputs)

    if success:
        return processed_df, results
    return None, None


async def process_remaining(df, selected_columns, test_size):
    """Process remaining samples after test batch."""
    chain = get_chain()
    if chain is None:
        return None

    # Validate inputs
    if not selected_columns:
        st.error("No column selected for student work")
        return None

    try:
        batch_df, inputs = prepare_batch(
            df,
            selected_columns,
            start_idx=test_size,
            batch_size=len(df) - test_size,  # Process all remaining
        )
        processed_df, results, success = await process_batch(chain, batch_df, inputs)

        if success:
            # Create a new dataframe with all columns
            full_df = df.copy()

            # Add new columns from results if they don't exist
            for col in processed_df.columns:
                if col not in full_df.columns:
                    full_df[col] = None

            # Update processed rows
            full_df.loc[processed_df.index] = processed_df
            return full_df

    except Exception as e:
        st.error(f"Error processing responses: {str(e)}")

    return None


def read_file_upload(uploaded_file):
    if uploaded_file is not None:
        # Get file extension
        file_ext = uploaded_file.name.split(".")[-1].lower()

        if file_ext == "docx":
            # Read DOCX files
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            content = []

            for element in doc.element.body:
                if element.tag.endswith("p"):  # Paragraph
                    paragraph = docx.Document(
                        io.BytesIO(uploaded_file.getvalue())
                    ).paragraphs[len(content)]
                    if paragraph.text.strip():
                        content.append(paragraph.text)

                elif element.tag.endswith("tbl"):  # Table
                    table = doc.tables[
                        len(
                            [
                                e
                                for e in doc.element.body[
                                    : doc.element.body.index(element)
                                ]
                                if e.tag.endswith("tbl")
                            ]
                        )
                    ]
                    table_data = []

                    # Get headers
                    headers = []
                    for cell in table.rows[0].cells:
                        headers.append(cell.text.strip())

                    # Get data rows
                    for row in table.rows[1:]:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        if any(cell for cell in row_data):  # Skip empty rows
                            table_data.append(dict(zip(headers, row_data)))

                    # Format table as text
                    if table_data:
                        content.append("\nTable:")
                        content.append(" | ".join(headers))
                        content.append(
                            "-"
                            * (sum(len(h) for h in headers) + 3 * (len(headers) - 1))
                        )
                        for row in table_data:
                            content.append(" | ".join(row.values()))
                        content.append("")

            return "\n".join(content)
        else:
            # Read TXT and MD files
            return uploaded_file.getvalue().decode("utf-8")
    return ""


async def handle_test_batch(df, selected_columns, test_size):
    """Handle the test batch processing."""
    test_df, results = await generate_feedback(df, selected_columns)

    if test_df is not None and results:
        # Create column config dynamically
        column_config = {}
        for col in test_df.columns:
            if col in selected_columns:
                continue
            if test_df[col].dtype == "object":
                column_config[col] = st.column_config.TextColumn(
                    col.replace("_", " ").title(), width="large"
                )
            else:
                column_config[col] = st.column_config.Column(
                    col.replace("_", " ").title()
                )

        # Store results in session state
        st.session_state.processed_df = test_df
        st.session_state.column_config = column_config
        return column_config
    return None


async def handle_remaining_samples(df, selected_columns, test_size, column_config):
    """Handle processing the remaining samples."""
    remaining_df = await process_remaining(df, selected_columns, test_size)
    if remaining_df is not None:
        st.write("### All Results")
        st.dataframe(remaining_df, column_config=column_config, hide_index=True)

        # Download all results
        csv_all = remaining_df.to_csv(index=False)
        st.download_button(
            label="Download All Results",
            data=csv_all,
            file_name="all_results.csv",
            mime="text/csv",
        )

        # Save full results to session state
        st.session_state.full_results_df = remaining_df


async def process_full_batch(df, selected_columns):
    """Process the full batch of responses."""
    # Process a small batch first to get the column structure
    test_df, results = await generate_feedback(df, selected_columns)

    if test_df is not None and results:
        # Create column config based on the test results
        column_config = {}
        for col in test_df.columns:
            if col in selected_columns:
                continue
            if test_df[col].dtype == "object":
                column_config[col] = st.column_config.TextColumn(
                    col.replace("_", " ").title(), width="large"
                )
            else:
                column_config[col] = st.column_config.Column(
                    col.replace("_", " ").title()
                )

        # Process all responses with the correct column config
        full_df = await process_remaining(
            df, selected_columns, 0
        )  # Start from beginning

        if full_df is not None:
            st.write("### All Results")
            st.dataframe(full_df, column_config=column_config, hide_index=True)

            # Download all results
            csv_all = full_df.to_csv(index=False)
            st.download_button(
                label="Download All Results",
                data=csv_all,
                file_name="all_results.csv",
                mime="text/csv",
            )

            # Save full results to session state
            st.session_state.full_results_df = full_df
            return True

    return False


def get_feedback_filename(original_filename=None, suffix=""):
    """Generate a feedback filename based on the original file."""
    # Use default name if no filename provided
    if not original_filename:
        base_name = "sample_responses"
    else:
        # Remove .csv extension if present
        base_name = original_filename.rsplit(".", 1)[0]
    # Add feedback prefix and suffix
    return f"Feedback - {base_name}{suffix}.csv"


# Instructions setup section
st.header("1. Instructions")

# Create two columns for side-by-side layout
col1, col2 = st.columns(2)

with col1:
    # Instructions input
    instructions = st.text_area(
        "Assessment Instructions and Rubric",
        value=DEFAULT_INSTRUCTIONS,
        height=600,
        help="Enter or paste the complete instructions including student instructions and grading rubric",
    )

with col2:
    with st.expander("Instructions Preview", expanded=True):
        if instructions:
            st.write(instructions)
        else:
            st.info("Enter instructions in the text area to see preview")

if instructions:
    st.session_state.instructions = instructions

# Process responses section
if not st.session_state.get("instructions"):
    st.stop()
else:
    st.header("2. Process Student Responses")

    # Load default data
    default_data_path = os.path.join("data", "transcriptions.csv")

    # File upload with default data option
    use_default = st.checkbox(
        "Use sample dataset",
        value=True,
        help="Use the provided sample dataset of student responses about peat and permafrost",
    )

    if use_default:
        try:
            df = pd.read_csv(default_data_path)
            st.success(f"Loaded {len(df)} sample responses")
        except Exception as e:
            st.error(f"Error loading default data: {str(e)}")
            st.stop()
    else:
        responses_file = st.file_uploader(
            "Upload Student Responses (CSV)", type=["csv"]
        )
        if responses_file is not None:
            df = pd.read_csv(responses_file)
            st.write(f"Loaded {len(df)} responses")
        else:
            st.info(
                "Please upload a CSV file with student responses or use the sample dataset"
            )
            st.stop()

    if df is not None:
        # Column selection
        st.write("### Select Response Columns")

        # Get previous selection if it exists and is valid
        previous_selection = st.session_state.previous_selection
        default_selection = [col for col in previous_selection if col in df.columns]

        # For default data, preselect the 'content' column
        if use_default and "content" in df.columns and not default_selection:
            default_selection = ["content"]

        selected_columns = st.multiselect(
            "Select columns containing student responses",
            df.columns,
            default=default_selection,
            help="Choose one or more columns that contain student work to be graded. Multiple columns will be combined.",
        )

        if selected_columns:
            st.session_state.selected_columns = selected_columns
            st.session_state.previous_selection = selected_columns

            # Show data preview of selected columns
            st.write("### Student Responses Preview")
            preview_df = df[selected_columns].head()
            st.dataframe(preview_df, hide_index=True)

            if len(df) > 5:
                st.info(f"Showing first 5 of {len(df)} responses")

            # Step 3: Process Samples
            st.subheader("3. Process Samples")

            # Sample size selection at the top
            test_size = st.number_input(
                "Sample size",
                min_value=1,
                max_value=len(df),
                value=1,
            )

            # Random seed for consistent results within a session
            if "random_seed" not in st.session_state:
                st.session_state.random_seed = np.random.randint(0, 10000)

            # Generate preview sample
            np.random.seed(st.session_state.random_seed)
            preview_indices = np.random.choice(len(df), size=test_size, replace=False)
            preview_sample = df.iloc[preview_indices]

            # Create two columns for side-by-side view
            col1, col2 = st.columns(2)

            with col1:
                st.markdown("#### Student Response")
                # Format each sample for better readability
                for idx, row in preview_sample.iterrows():
                    st.markdown("---")  # Separator between samples if more than one

                    # Show metadata if available
                    if "filename" in row:
                        st.caption(f"Source: {row['filename']}")
                    if "prompt" in row:
                        st.markdown(f"**Prompt:** {row['prompt']}")

                    # Show the actual response with proper formatting
                    for col in selected_columns:
                        if col not in ["filename", "prompt"]:  # Skip metadata columns
                            content = str(row[col])
                            # Format the content with proper paragraphs
                            paragraphs = content.split("\n")
                            formatted_paragraphs = []
                            for p in paragraphs:
                                if p.strip():  # Skip empty paragraphs
                                    # Wrap text at 80 characters
                                    wrapped = fill(
                                        p.strip(),
                                        width=80,
                                        break_long_words=False,
                                        break_on_hyphens=False,
                                    )
                                    formatted_paragraphs.append(wrapped)

                            st.markdown("\n\n".join(formatted_paragraphs))

            with col2:
                st.markdown("#### Generated Feedback")
                if (
                    "processed_df" not in st.session_state
                    or st.session_state.processed_df is None
                ):
                    # Show the generate feedback button when no feedback exists
                    if st.button("Generate Student Feedback", type="primary"):
                        with st.spinner("Generating feedback..."):
                            try:
                                df_sample = preview_sample.copy()
                                df = df_sample  # Use the random sample for processing
                                st.session_state.test_size = test_size

                                # Process the sample
                                column_config = asyncio.run(
                                    handle_test_batch(
                                        df, st.session_state.selected_columns, test_size
                                    )
                                )

                                if column_config is not None:
                                    st.success("Feedback generated successfully!")
                                else:
                                    st.error("Failed to process sample")
                            except Exception as e:
                                st.error(f"Error processing sample: {str(e)}")
                else:
                    # Show the feedback if it exists
                    processed_df = st.session_state.processed_df
                    for idx, row in processed_df.iterrows():
                        st.markdown("---")

                        # Display strengths and improvement sections
                        if "strengths" in row:
                            st.markdown("##### 💪 Strengths")
                            content = str(row["strengths"])
                            paragraphs = content.split("\n")
                            formatted_paragraphs = []
                            for p in paragraphs:
                                if p.strip():
                                    wrapped = fill(
                                        p.strip(),
                                        width=80,
                                        break_long_words=False,
                                        break_on_hyphens=False,
                                    )
                                    formatted_paragraphs.append(wrapped)
                            st.markdown("\n\n".join(formatted_paragraphs))
                            st.markdown("")  # Add spacing

                        if "improvement" in row:
                            st.markdown("##### 🎯 Areas for Improvement")
                            content = str(row["improvement"])
                            paragraphs = content.split("\n")
                            formatted_paragraphs = []
                            for p in paragraphs:
                                if p.strip():
                                    wrapped = fill(
                                        p.strip(),
                                        width=80,
                                        break_long_words=False,
                                        break_on_hyphens=False,
                                    )
                                    formatted_paragraphs.append(wrapped)
                            st.markdown("\n\n".join(formatted_paragraphs))
                            st.markdown("")  # Add spacing

                        # Display any additional feedback fields that might be present
                        for col in processed_df.columns:
                            if (
                                col not in ["strengths", "improvement"]
                                and col not in selected_columns
                                and col not in ["filename", "prompt"]
                            ):
                                st.markdown(f"##### {col.replace('_', ' ').title()}")
                                content = str(row[col])
                                paragraphs = content.split("\n")
                                formatted_paragraphs = []
                                for p in paragraphs:
                                    if p.strip():
                                        wrapped = fill(
                                            p.strip(),
                                            width=80,
                                            break_long_words=False,
                                            break_on_hyphens=False,
                                        )
                                        formatted_paragraphs.append(wrapped)
                                st.markdown("\n\n".join(formatted_paragraphs))
                                st.markdown("")  # Add spacing between sections

                    # Add a button to regenerate feedback
                    if st.button("Regenerate Feedback", type="primary"):
                        with st.spinner("Generating feedback..."):
                            try:
                                df_sample = preview_sample.copy()
                                df = df_sample  # Use the random sample for processing
                                st.session_state.test_size = test_size
                                st.session_state.processed_df = None

                                # Process the sample
                                column_config = asyncio.run(
                                    handle_test_batch(
                                        df, st.session_state.selected_columns, test_size
                                    )
                                )

                                if column_config is not None:
                                    st.success("Feedback regenerated successfully!")
                                else:
                                    st.error("Failed to process sample")
                            except Exception as e:
                                st.error(f"Error processing sample: {str(e)}")
                            st.rerun()

            # Process all responses option at the bottom
            st.markdown("---")
            st.markdown("#### Process All Responses")
            col3, col4 = st.columns([3, 1])
            with col3:
                st.write(f"Generate feedback for all {len(df)} responses")
            with col4:
                if st.button("Process All", type="primary"):
                    with st.spinner(f"Processing {len(df)} responses..."):
                        try:
                            success = asyncio.run(
                                process_full_batch(
                                    df, st.session_state.selected_columns
                                )
                            )
                            if success:
                                st.success("Processing complete!")
                            else:
                                st.error("Failed to process responses")
                        except Exception as e:
                            st.error(f"Error processing responses: {str(e)}")

            # Show download buttons if results exist
            if (
                "processed_df" in st.session_state
                and st.session_state.processed_df is not None
            ):
                csv_test = st.session_state.processed_df.to_csv(index=False)
                feedback_filename = get_feedback_filename(
                    responses_file.name if not use_default and responses_file else None,
                    f" - {test_size} samples",
                )
                st.download_button(
                    label="Download Results",
                    data=csv_test,
                    file_name=feedback_filename,
                    mime="text/csv",
                )

            if st.session_state.get("full_results_df") is not None:
                feedback_filename = get_feedback_filename(
                    responses_file.name if not use_default and responses_file else None
                )
                csv_all = st.session_state.full_results_df.to_csv(index=False)
                st.download_button(
                    label="Download All Results",
                    data=csv_all,
                    file_name=feedback_filename,
                    mime="text/csv",
                )
